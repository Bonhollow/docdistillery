import json
import os
from typing import Any, Dict, Optional


def export_summary(summary_obj: Dict[str, Any], out_path: Optional[str] = None, format: str = "md") -> str:
    """
    Exports a summary object to various formats.

    Args:
        summary_obj (Dict): The summary data (tldr, executive, sections, provenance, metadata, doc_id).
        out_path (Optional[str]): The file path to write to. If None, returns string (for text formats).
        format (str): "md", "txt", "json", "docx", "pdf".

    Returns:
        str: The formatted content or the output path.

    Raises:
        ValueError: If out_path is required but not provided.
        ImportError: If an optional dependency is missing for the requested format.
    """
    format = format.lower()
    content = ""

    # 1. Formatting logic
    if format == "json":
        content = json.dumps(summary_obj, indent=2)
    elif format == "md":
        title = summary_obj.get("metadata", {}).get("filename", "Summary")
        content = f"# {title}\n\n"
        content += f"## TL;DR\n{summary_obj.get('tldr', '')}\n\n"

        exec_sum = summary_obj.get("executive", [])
        if exec_sum:
            content += "## Executive Summary\n"
            for point in exec_sum:
                content += f"- {point}\n"
            content += "\n"

        for section in summary_obj.get("sections", []):
            content += f"### {section.get('title', 'Section')}\n"
            content += f"{section.get('content', '')}\n\n"

    elif format == "txt":
        # Simplified text format
        title = summary_obj.get("metadata", {}).get("filename", "Summary")
        content = f"{title}\n" + "=" * len(title) + "\n\n"
        content += f"TL;DR: {summary_obj.get('tldr', '')}\n\n"
        for point in summary_obj.get("executive", []):
            content += f"* {point}\n"
        content += "\n"
        for section in summary_obj.get("sections", []):
            content += f"{section.get('title', 'SECTION')}\n"
            content += f"{section.get('content', '')}\n\n"

    elif format in ["docx", "pdf"]:
        if out_path is None:
            raise ValueError(f"Format '{format}' requires out_path.")

        if format == "docx":
            try:
                from docx import Document
            except ImportError as err:
                raise ImportError(
                    "Format 'docx' requires 'python-docx'. Install it with 'pip install python-docx'."
                ) from err

            doc = Document()
            title = summary_obj.get("metadata", {}).get("filename", "Summary")
            doc.add_heading(title, 0)

            doc.add_heading("TL;DR", level=1)
            doc.add_paragraph(summary_obj.get("tldr", ""))

            if summary_obj.get("executive"):
                doc.add_heading("Executive Summary", level=1)
                for point in summary_obj.get("executive", []):
                    doc.add_paragraph(point, style="List Bullet")

            for section in summary_obj.get("sections", []):
                doc.add_heading(section.get("title", "Section"), level=2)
                doc.add_paragraph(section.get("content", ""))

            doc.save(out_path)

        elif format == "pdf":
            try:
                from weasyprint import HTML
            except (ImportError, OSError) as err:
                msg = "Format 'pdf' requires 'weasyprint' and its system dependencies (pango, cairo, libffi)."
                if "libgobject" in str(err) or "cannot load library" in str(err).lower():
                    msg += " On macOS, please install them via Homebrew: 'brew install pango'."
                raise ImportError(msg) from err

            # Use the markdown representation as a base for simple conversion to HTML
            md_content = export_summary(summary_obj, out_path=None, format="md")
            # Very basic markdown-to-html conversion for the sake of this example
            html_body = (
                md_content.replace("# ", "<h1>").replace("## ", "<h2>").replace("### ", "<h3>").replace("\n", "<br>")
            )
            html_content = f"<html><body>{html_body}</body></html>"
            try:
                HTML(string=html_content).write_pdf(out_path)
            except Exception as err:
                msg = f"PDF conversion failed: {err}"
                if "cannot load library" in str(err).lower() or "libgobject" in str(err).lower():
                    msg += " This is likely due to missing system dependencies. On macOS, run: 'brew install pango'."
                raise RuntimeError(msg) from err
    else:
        raise ValueError(f"Unsupported format: {format}")

    # 2. File writing and Provenance
    if out_path and format not in ["docx", "pdf"]:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(content)

    if out_path:
        doc_id = summary_obj.get("doc_id", "unknown")
        dir_name = os.path.dirname(out_path)
        prov_path = os.path.join(dir_name, f"provenance_{doc_id}.json")
        with open(prov_path, "w", encoding="utf-8") as f:
            json.dump(summary_obj.get("provenance", {}), f, indent=2)
        return out_path

    return content
